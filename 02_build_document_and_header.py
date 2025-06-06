import os
import pandas as pd
import re
import shutil
import datetime
from docx import Document
from docxcompose.composer import Composer  # Add this import
import docx.shared
import docxcompose.composer as composer
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Define paths
current_dir = os.path.dirname(os.path.abspath(__file__))
data_folder = os.path.join(current_dir, 'data')
template_file = os.path.join(data_folder, 'SAA-DBL-TranscriptionTemplate.docx')
input_data_file = os.path.join(current_dir, 'temp', 'processed_data.pkl')

# Define output directory for generated documents
OUTPUT_DIR = os.path.join(current_dir, "generated_documents")
os.makedirs(OUTPUT_DIR, exist_ok=True)

def format_citation_text(row, content_type):
    """
    Format citation components for use with formatted insertion
    Returns a dictionary with different components of the citation
    """
    doc_number = str(row.get('DBL - Doc number', ''))
    date_value = str(row.get('Date', ''))
    volume = str(row.get('volume', '')).strip()
    
    # Use correct capitalization for range columns
    transcript_range = str(row.get('Transcript range', ''))
    translate_range = str(row.get('Translate range', ''))
    
    # Format differs based on document type
    if content_type == 'Transcript':
        doc_type = "transcription"
        page_range = transcript_range
    else:  # Translation
        doc_type = "translation"
        page_range = translate_range
    
    # Determine the book title and additional info based on volume
    if volume == 'I':
        book_title = "Documents of Brotherly Love: Dutch Mennonite Aid to Swiss Anabaptists"
        volume_info = "Volume 1, 1635-1709"
        editors = "edited by David J. Rempel Smucker and John L. Ruth"
        publisher = "(Millersburg, OH: Ohio Amish Library, 2007)"
    elif volume == 'II':
        book_title = "Documents of Brotherly Love: Dutch Mennonite Aid to Swiss Anabaptists"
        volume_info = "Volume II, 1710-1711"
        editors = ""
        publisher = "(Millersburg, OH: Ohio Amish Library, 2015)"
    else:
        # Default case
        book_title = "Documents of Brotherly Love: Dutch Mennonite Aid to Swiss Anabaptists"
        volume_info = ""
        editors = ""
        publisher = "(Millersburg, OH: Ohio Amish Library)"
    
    # Return components separately
    return {
        "author": "James W. Lowry",
        "doc_number": doc_number,
        "date": date_value,
        "doc_type": doc_type,
        "book_title": book_title,
        "volume_info": volume_info,
        "editors": editors,
        "publisher": publisher,
        "page_range": page_range
    }

def add_formatted_citation(doc, citation_components):
    """Add citation with proper formatting"""
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().startswith("Citation:"):
            # Found the citation paragraph
            print(f"Found Citation line: '{para.text}'")
            
            # Clear the paragraph
            for run in para.runs:
                run.clear()
            
            # Add the citation parts with appropriate formatting
            # Part 1: "Citation: " label in bold
            run1 = para.add_run("Citation: ")
            run1.bold = True
            
            # Part 2: Author and beginning of citation
            para.add_run(f"{citation_components['author']}, \"Document {citation_components['doc_number']}, ")
            para.add_run(f"{citation_components['date']}, {citation_components['doc_type']},\" in ")
            
            # Part 3: Book title in italics
            run_title = para.add_run(citation_components['book_title'])
            run_title.italic = True
            
            # Part 4: Volume info and the rest
            if citation_components['volume_info']:
                para.add_run(f" {citation_components['volume_info']}")
            
            if citation_components['editors']:
                para.add_run(f", {citation_components['editors']}")
            
            para.add_run(f" {citation_components['publisher']}, {citation_components['page_range']}.")
            
            print(f"Added formatted citation with italicized book title")
            return True
    
    print("Warning: Citation line not found in document")
    return False

def update_digital_id_in_header(header, new_digital_id):
    """Update the Digital ID in a header, checking both paragraphs and tables"""
    # First check header paragraphs
    for para in header.paragraphs:
        if "Digital ID:" in para.text:
            print(f"Found Digital ID in header paragraph: '{para.text}'")
            # Clear the paragraph
            for run in para.runs:
                run.clear()
            # Create new text with proper formatting
            run1 = para.add_run("Digital ID: ")
            run1.bold = True
            para.add_run(new_digital_id)
            print(f"Updated paragraph text to: 'Digital ID: {new_digital_id}'")
            return True
    
    # If not found in paragraphs, check header tables
    for table in header.tables:
        print(f"Examining header table with {len(table.rows)} rows")
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if "Digital ID:" in para.text:
                        print(f"Found Digital ID in table cell: '{para.text}'")
                        # Clear the paragraph
                        for run in para.runs:
                            run.clear()
                        # Create new text with proper formatting
                        run1 = para.add_run("Digital ID: ")
                        run1.bold = True
                        para.add_run(new_digital_id)
                        print(f"Updated cell text to: 'Digital ID: {new_digital_id}'")
                        return True
    
    print("WARNING: Digital ID not found in header paragraphs or tables")
    return False

def add_metadata_fields(doc, row, content_type):
    """
    Update metadata fields in the document:
    - Copyright: Leave as is in template
    - Other fields: Add tabs after ": " and add value, or delete line if no value exists
    
    Tab rules:
    - 5 or fewer characters: 3 tabs
    - 6-11 characters: 2 tabs
    - 12+ characters: 1 tab
    """
    print("Processing metadata fields...")
    
    # Define metadata fields to process (excluding Copyright)
    metadata_fields = [
        {"tag": "Date", "column": "Date"},
        {"tag": "Sender", "column": "Sender"},
        {"tag": "Sender Place", "column": "Sender Place"},
        {"tag": "Receiver", "column": "Receiver"},
        {"tag": "Receiver Place", "column": "Receiver Place"}
    ]
    
    # Handle Language field (might differ based on content type)
    if content_type == 'Translate':
        language_value = "English"
    else:
        language_value = row.get('Language', '')
    
    metadata_fields.append({"tag": "Language", "column": None, "value": language_value})
    
    # Find paragraphs that begin with each metadata tag
    paragraphs_to_delete = []
    
    for i, para in enumerate(doc.paragraphs):
        # Skip empty paragraphs
        if not para.text.strip():
            continue
            
        # Skip Copyright paragraph - leave it as is
        if para.text.strip().startswith("Copyright:"):
            print(f"Keeping original Copyright text: '{para.text}'")
            continue
        
        # Check for metadata fields
        for field in metadata_fields:
            if para.text.strip().startswith(f"{field['tag']}:"):
                print(f"Found {field['tag']} field: '{para.text}'")
                
                # Get the value from dataframe or from pre-defined value
                if 'value' in field:
                    value = field['value']
                else:
                    value = row.get(field['column'], '')
                
                # If the value is None, NaN, or empty string, mark paragraph for deletion
                if pd.isna(value) or str(value).strip() == '':
                    print(f"  No value or NaN for {field['tag']}, will remove line")
                    paragraphs_to_delete.append(i)
                else:
                    # Clear the paragraph and rebuild it
                    for run in para.runs:
                        run.clear()
                    
                    # Add tag with bold formatting
                    bold_run = para.add_run(f"{field['tag']}: ")
                    bold_run.bold = True
                    
                    # Add tabs based on tag length
                    if len(field['tag']) <= 5:
                        para.add_run("\t\t\t")  # 3 tabs for 5 or fewer chars
                    elif len(field['tag']) < 12:
                        para.add_run("\t\t")     # 2 tabs for 6-11 chars
                    else:
                        para.add_run("\t")       # 1 tab for 12+ chars
                    
                    # Add value
                    para.add_run(str(value))
                    print(f"  Updated with value: '{value}'")
                
                break
    
    # Delete paragraphs marked for deletion (in reverse order to maintain indices)
    for idx in sorted(paragraphs_to_delete, reverse=True):
        p = doc.paragraphs[idx]._p
        p.getparent().remove(p)
        print(f"Removed paragraph at index {idx}")
    
    return doc

def remove_instruction_text(doc):
    """
    Remove the specific instruction text from the document
    """
    instruction_text = "<For the following PRINT fields, if they aren't available for a document, remove them."
    
    print("Searching for instruction text...")
    
    # Find paragraphs containing any part of the instruction text (more robust)
    paragraphs_to_delete = []
    for i, para in enumerate(doc.paragraphs):
        # Check for the start of the instruction text to identify the paragraph
        if instruction_text in para.text:
            print(f"Found instruction text at paragraph {i}: '{para.text[:50]}...'")
            paragraphs_to_delete.append(i)
        # Also check for sections of text in case it's broken up
        elif "<For the following PRINT fields" in para.text or "remove them. For example" in para.text:
            print(f"Found partial instruction text at paragraph {i}: '{para.text[:50]}...'")
            paragraphs_to_delete.append(i)
    
    # Delete paragraphs with instruction text (in reverse order to maintain indices)
    if paragraphs_to_delete:
        for idx in sorted(paragraphs_to_delete, reverse=True):
            p = doc.paragraphs[idx]._p
            p.getparent().remove(p)
            print(f"Removed instruction paragraph at index {idx}")
        print(f"Removed {len(paragraphs_to_delete)} paragraphs containing instruction text")
    else:
        print("No instruction text found")
    
    return doc

def copy_content_from_source(doc, row, content_type, data_folder, output_filename):
    """
    Copy content from source file while preserving footnotes using docxcompose
    """
    print("\n==== CONTENT COPYING WITH FOOTNOTES ====")
    print(f"Document being created: {output_filename}")
    print(f"Content type: {content_type}")
    
    # Determine which column to use based on content_type
    source_filename = row.get(content_type)
    
    # If no source filename provided, nothing to copy
    if pd.isna(source_filename) or not source_filename:
        print(f"No {content_type} source file specified for this document")
        return doc
    
    # Build full path to source file
    source_path = os.path.join(data_folder, 'transcriptions-translations', source_filename)
    
    # Check if source file exists
    if not os.path.exists(source_path):
        print(f"WARNING: Source file not found: {source_path}")
        return doc
    
    print(f"Source document: {source_path}")
    
    try:
        import tempfile
        
        # Create temporary files
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            temp_path = temp_file.name
        
        # Save the current document state to temp file
        doc.save(temp_path)
        
        # Create another temporary file for the composed result
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as result_file:
            result_path = result_file.name
            
        # Remove the field we don't want
        target_field = "Translation:" if content_type == "Translate" else "Transcription:"
        other_field = "Transcription:" if content_type == "Translate" else "Translation:"
        
        paragraphs_to_delete = []
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip().startswith(other_field):
                paragraphs_to_delete.append(i)
                print(f"Removing {other_field} field")
        
        for idx in sorted(paragraphs_to_delete, reverse=True):
            p = doc.paragraphs[idx]._p
            p.getparent().remove(p)
        
        # Make sure the target field exists
        field_found = False
        for para in doc.paragraphs:
            if para.text.strip().startswith(target_field):
                field_found = True
                break
                
        if not field_found:
            print(f"Adding {target_field} field")
            field_para = doc.add_paragraph()
            bold_run = field_para.add_run(f"{target_field} ")
            bold_run.bold = True
        
        # Save the modified document
        doc.save(temp_path)
        
        # Load the source document with docxcompose
        master = Document(temp_path)
        composer = Composer(master)
        
        # Load source document
        source_doc = Document(source_path)
        
        # Append the content (this preserves footnotes)
        composer.append(source_doc)
        
        # Save the composed document
        composer.save(result_path)
        
        # Load the resulting document
        result_doc = Document(result_path)
        
        # Clean up temporary files
        try:
            os.unlink(temp_path)
            os.unlink(result_path)
        except:
            pass
        
        print(f"Successfully copied content with footnotes from {source_filename}")
        return result_doc
        
    except Exception as e:
        import traceback
        print(f"Error copying content with footnotes: {str(e)}")
        print(traceback.format_exc())
        
        # Add error message to document
        error_para = doc.add_paragraph()
        error_run = error_para.add_run(f"[ERROR COPYING CONTENT: {str(e)}]")
        error_run.bold = True
        
        return doc
def create_document(row, content_type):
    """Create a new document based on the template and row data"""
    # Skip if the specified column has no value
    if pd.isna(row[content_type]):
        return None
    
    # Get values for filename
    filename_value = str(row.get('Filename', ''))
    
    # Determine language code for the filename
    if content_type == 'Translate':
        # Always use EN for translations
        language_code = 'EN'
    else:
        # For transcriptions, map language names to codes
        language_value = str(row.get('Language', ''))
        
        # Check for combined language cases
        if language_value.lower() == 'german/dutch':
            language_code = 'DE-NL'
        elif language_value.lower() == 'french; german':
            language_code = 'FR-DE'
        elif language_value.lower() == 'german':
            language_code = 'DE'
        elif language_value.lower() == 'dutch':
            language_code = 'NL'
        elif language_value.lower() == 'french':
            language_code = 'FR'
        else:
            # Keep original value for other languages
            language_code = language_value
    
    # Create a filename using the specified format with appropriate language code
    safe_filename = f"{filename_value}_{language_code}_{content_type.lower()}.docx"
    # Clean up filename to be valid (remove illegal characters)
    safe_filename = re.sub(r'[<>:"/\\|?*]', '', safe_filename)
    # Replace spaces with underscores
    safe_filename = safe_filename.replace(" ", "_")
    
    output_path = os.path.join(OUTPUT_DIR, safe_filename)
    
    # Digital ID and other values still needed for document content
    digital_id = str(row.get('Digital ID', 'unknown'))
    doc_number = str(row.get('DBL - Doc number', 'unknown'))
    date_value = str(row.get('Date', 'unknown'))
    
    # Create a copy of the template
    print(f"Creating document: {output_path}")
    shutil.copy(template_file, output_path)
    
    # Open the document
    doc = Document(output_path)
    
    # Remove instruction text from the document
    doc = remove_instruction_text(doc)
    
    # NEW: Remove placeholder text and Document Type lines
    paragraphs_to_delete = []
    for i, para in enumerate(doc.paragraphs):
        # Check for placeholder text
        if "<copy/paste transcription here>" in para.text:
            print(f"Found placeholder text to remove at paragraph {i}")
            paragraphs_to_delete.append(i)
        
        # Check for Document Type line
        if para.text.strip().startswith("Document Type:"):
            print(f"Found Document Type line to remove at paragraph {i}")
            paragraphs_to_delete.append(i)
    
    # Delete marked paragraphs (in reverse order to maintain indices)
    for idx in sorted(paragraphs_to_delete, reverse=True):
        p = doc.paragraphs[idx]._p
        p.getparent().remove(p)
        print(f"Removed paragraph at index {idx}")
    
    # Process the header sections
    for section_idx, section in enumerate(doc.sections):
        header = section.header
        update_digital_id_in_header(header, digital_id)
    
    # Generate the enhanced citation text
    citation_components = format_citation_text(row, content_type)
    
    # Add formatted citation to the document
    add_formatted_citation(doc, citation_components)
    
    # Process metadata fields
    doc = add_metadata_fields(doc, row, content_type)
    
    # Add source content based on content type - now with proper footnote handling
    doc = copy_content_from_source(doc, row, content_type, data_folder, os.path.basename(output_path))
    
    # Save the modified document
    doc.save(output_path)
    
    return output_path

def main():
    # Check if the processed data file exists
    if not os.path.exists(input_data_file):
        print(f"Error: Processed data file not found: {input_data_file}")
        print("Please run step 1 first (make run)")
        return
    
    # Load the processed data
    try:
        df = pd.read_pickle(input_data_file)
        print(f"Loaded data with {len(df)} rows from {input_data_file}")
    except Exception as e:
        print(f"Error loading processed data: {e}")
        return
    
    # Track created documents
    transcript_docs = []
    translate_docs = []
    
    # Process each row in the dataframe
    for idx, row in df.iterrows():
        # Create transcript document if column has value
        if 'Transcript' in df.columns and not pd.isna(row.get('Transcript', pd.NA)):
            doc_path = create_document(row, 'Transcript')
            if doc_path:
                transcript_docs.append(doc_path)
                print(f"Created Transcript document: {os.path.basename(doc_path)}")
        
        # Create translate document if column has value
        if 'Translate' in df.columns and not pd.isna(row.get('Translate', pd.NA)):
            doc_path = create_document(row, 'Translate')
            if doc_path:
                translate_docs.append(doc_path)
                print(f"Created Translate document: {os.path.basename(doc_path)}")
    
    print(f"\nGenerated {len(transcript_docs)} Transcript documents")
    print(f"Generated {len(translate_docs)} Translate documents")
    print(f"Documents saved to: {os.path.abspath(OUTPUT_DIR)}")

if __name__ == "__main__":
    main()
