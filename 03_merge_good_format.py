import os
import pandas as pd
import shutil
from datetime import datetime
from docx import Document
from docxcompose.composer import Composer

def get_lang_code(language, content_type):
    if content_type == 'translate':
        return 'EN'
    l = language.lower()
    if l == 'german/dutch':
        return 'DE-NL'
    if l == 'french; german':
        return 'FR-DE'
    if l == 'german':
        return 'DE'
    if l == 'dutch':
        return 'NL'
    if l == 'french':
        return 'FR'
    return language

def replace_section_with_footnotes(target_path, footnote_path, section_label):
    # Save to a temp file, then use Composer to append footnote content
    temp_path = target_path + ".tmp"
    shutil.copy(target_path, temp_path)
    target_doc = Document(temp_path)
    composer = Composer(target_doc)
    footnote_doc = Document(footnote_path)

    # Find the section label (e.g., "Transcription:" or "Translation:") in the target doc
    section_idx = None
    for i, para in enumerate(target_doc.paragraphs):
        if para.text.strip().startswith(section_label):
            section_idx = i
            break
    if section_idx is None:
        print(f"Section label '{section_label}' not found in {target_path}")
        os.remove(temp_path)
        return False

    # Remove all paragraphs after the section label up to the next metadata field or end
    # (Assume metadata fields are bold and end with ':')
    i = section_idx + 1
    while i < len(target_doc.paragraphs):
        para = target_doc.paragraphs[i]
        if any(run.bold and run.text.strip().endswith(':') for run in para.runs):
            break
        p = para._p
        p.getparent().remove(p)
        # Do not increment i, as the list shrinks

    # Save the truncated document
    target_doc.save(temp_path)

    # Now append the footnote_doc content after the section label
    # (Composer appends at the end, so we need to move the section label to the end, append, then move back)
    # Instead, we will use Word's structure: copy the section label paragraph, clear after, append, then save
    # This is a workaround for docx limitations

    # Open again for appending
    target_doc = Document(temp_path)
    composer = Composer(target_doc)
    composer.append(footnote_doc)
    composer.save(temp_path)

    # Now, move the appended content to just after the section label
    # Reload and manipulate paragraphs
    doc = Document(temp_path)
    section_idx = None
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().startswith(section_label):
            section_idx = i
            break
    if section_idx is None:
        print(f"Section label '{section_label}' not found after append in {target_path}")
        os.remove(temp_path)
        return False

    # Find where the appended content starts (first non-empty after section_idx)
    appended_start = section_idx + 1
    while appended_start < len(doc.paragraphs) and not doc.paragraphs[appended_start].text.strip():
        appended_start += 1

    # Move appended paragraphs to just after section label
    appended_paragraphs = []
    i = appended_start
    while i < len(doc.paragraphs):
        appended_paragraphs.append(doc.paragraphs[i])
        i += 1
    # Remove them from the end
    for para in appended_paragraphs:
        p = para._p
        p.getparent().remove(p)
    # Insert them after section_idx
    for para in reversed(appended_paragraphs):
        doc._body._body.insert(section_idx + 1, para._p)

    # Save the final document
    doc.save(target_path)
    os.remove(temp_path)
    return True

def main():
    # Paths
    base_dir = os.path.dirname(os.path.abspath(__file__))
    data_dir = os.path.join(base_dir, "data")
    footnotes_dir = os.path.join(data_dir, "DBL-UpdatedFootnotes")
    generated_dir = os.path.join(base_dir, "generated_documents")
    excel_path = os.path.join(data_dir, "SAA-DBL-MergeData.xlsx")
    today_str = datetime.now().strftime("%Y%m%d")
    output_dir = os.path.join(base_dir, f"generated_docs_updated_{today_str}")
    os.makedirs(output_dir, exist_ok=True)

    # Read Excel
    df = pd.read_excel(excel_path)
    transcript_files = set(df['Transcript'].dropna().astype(str))
    translate_files = set(df['Translate'].dropna().astype(str))

    # List all .docx files in DBL-UpdatedFootnotes
    footnote_files = [f for f in os.listdir(footnotes_dir) if f.lower().endswith('.docx')]

    for fn_file in footnote_files:
        match_type = None
        if fn_file in transcript_files:
            match_type = 'Transcript'
        elif fn_file in translate_files:
            match_type = 'Translate'
        else:
            continue

        rows = df[df[match_type] == fn_file]
        for _, row in rows.iterrows():
            filename = str(row.get('Filename', ''))
            language = str(row.get('Language', ''))
            content_type = match_type.lower()
            lang_code = get_lang_code(language, content_type)
            gen_doc = f"{filename}_{lang_code}_{content_type}.docx".replace(" ", "_")
            gen_doc_path = os.path.join(generated_dir, gen_doc)
            if not os.path.exists(gen_doc_path):
                print(f"Generated document not found: {gen_doc_path}")
                continue

            # Copy to output directory
            out_doc_path = os.path.join(output_dir, gen_doc)
            shutil.copy(gen_doc_path, out_doc_path)

            # Replace the section with the updated footnote file
            section_label = "Transcription:" if match_type == "Transcript" else "Translation:"
            footnote_path = os.path.join(footnotes_dir, fn_file)
            print(f"Updating {out_doc_path} with {footnote_path} in section {section_label}")
            replace_section_with_footnotes(out_doc_path, footnote_path, section_label)

    print(f"All updates complete. Output in {output_dir}")

if __name__ == "__main__":
    main()