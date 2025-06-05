import os
from docx import Document
from docxcompose.composer import Composer
import shutil

def test_footnote_extraction():
    """Test extraction and insertion of text with footnotes"""
    # Define file paths
    source_path = r"C:\Users\ch738340\OneDrive - University of Central Florida\Documents\CHDR\PRINT Project\data exploration\SAA_doc_builder\data\transcriptions-translations\565.A_1013,1065_10-10-1710-Dut.docx"
    target_path = r"C:\Users\ch738340\OneDrive - University of Central Florida\Documents\CHDR\PRINT Project\data exploration\SAA_doc_builder\data\TEST DOCUMENT FOR INSERTION.docx"
    output_path = r"C:\Users\ch738340\OneDrive - University of Central Florida\Documents\CHDR\PRINT Project\data exploration\SAA_doc_builder\data\RESULT WITH FOOTNOTES.docx"
    
    # Make sure files exist
    if not os.path.exists(source_path):
        print(f"Source file not found: {source_path}")
        return
        
    if not os.path.exists(target_path):
        print(f"Target file not found: {target_path}")
        return
    
    print(f"Source: {source_path}")
    print(f"Target: {target_path}")
    print(f"Output: {output_path}")
    
    try:
        # Method 1: Using docxcompose (most reliable for footnotes)
        # First create a copy of the target so we don't modify the original
        shutil.copy(target_path, output_path)
        
        # Load the documents
        target_doc = Document(output_path)
        
        # Add a header to indicate where the inserted content begins
        target_doc.add_paragraph("=== INSERTED CONTENT BELOW ===").style = 'Heading1'
        target_doc.add_paragraph()  # Add space
        
        # Save the modified target document
        target_doc.save(output_path)
        
        # Now use docxcompose to append the source document
        master = Document(output_path)
        composer = Composer(master)
        source_doc = Document(source_path)
        composer.append(source_doc)
        
        # Save the composed document
        composer.save(output_path)
        
        print(f"Document with footnotes saved to {output_path}")
        
        # Method 2: Alternative approach using lxml if docxcompose doesn't work
        # This is a fallback and would require more complex XML manipulation
        
    except Exception as e:
        print(f"Error: {str(e)}")

if __name__ == "__main__":
    test_footnote_extraction()